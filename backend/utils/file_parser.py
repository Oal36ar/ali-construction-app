import fitz  # PyMuPDF
import pandas as pd
from docx import Document
import io
import mimetypes
from typing import Dict, Any, Tuple
import csv
from openpyxl import load_workbook

def parse_file_by_type(filename: str, content: bytes) -> Dict[str, str]:
    """
    Parse file content based on file type and return text + preview
    
    Args:
        filename: Name of the uploaded file
        content: Raw bytes content of the file
        
    Returns:
        dict with keys: "text" (extracted content), "preview" (summary info)
    """
    file_type = detect_file_type(filename, None)
    
    try:
        if file_type == "pdf":
            text, metadata = _parse_pdf_content(content, filename)
            preview = f"{metadata.get('pages', 0)} pages extracted"
        elif file_type == "csv":
            text, metadata = _parse_csv_content(content, filename)
            preview = f"{metadata.get('rows', 0)} rows, {len(metadata.get('columns', []))} columns"
        elif file_type == "xlsx":
            text, metadata = _parse_xlsx_content(content, filename)
            preview = f"{len(metadata.get('sheets', []))} sheets extracted"
        elif file_type == "docx":
            text, metadata = _parse_docx_content(content, filename)
            preview = f"{metadata.get('paragraphs', 0)} paragraphs, {metadata.get('tables', 0)} tables"
        else:
            # Default: decode first 1000 chars raw
            try:
                text = content.decode('utf-8')[:1000]
                preview = f"First 1000 characters extracted"
            except UnicodeDecodeError:
                text = content.decode('latin-1', errors='ignore')[:1000]
                preview = f"First 1000 characters extracted (latin-1)"
        
        return {"text": text, "preview": preview}
        
    except Exception as e:
        return {
            "text": f"Error parsing {filename}: {str(e)}",
            "preview": f"Failed to parse {file_type} file"
        }

def detect_file_type(filename: str, content_type: str = None) -> str:
    """Detect file type from filename and content type"""
    if content_type:
        if "pdf" in content_type:
            return "pdf"
        elif "csv" in content_type:
            return "csv"
        elif "excel" in content_type or "spreadsheet" in content_type:
            return "xlsx"
        elif "word" in content_type or "document" in content_type:
            return "docx"
        elif "text" in content_type:
            return "text"
    
    # Fallback to file extension
    extension = filename.lower().split('.')[-1] if '.' in filename else ""
    type_mapping = {
        'pdf': 'pdf',
        'csv': 'csv',
        'xlsx': 'xlsx',
        'xls': 'xlsx',
        'docx': 'docx',
        'doc': 'docx',
        'txt': 'text'
    }
    return type_mapping.get(extension, 'unknown')

def _parse_pdf_content(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
    """Parse PDF file using PyMuPDF"""
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        text_content = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_content.append(f"--- Page {page_num + 1} ---\n")
            text_content.append(page.get_text())
            text_content.append("\n")
        
        doc.close()
        
        full_text = "".join(text_content)
        metadata = {
            "pages": len(doc),
            "type": "pdf",
            "title": filename
        }
        
        return full_text, metadata
        
    except Exception as e:
        return f"Error parsing PDF: {str(e)}", {"error": str(e), "type": "pdf"}

def _parse_csv_content(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
    """Parse CSV file using pandas"""
    try:
        # Try UTF-8 first, then fallback encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                df = pd.read_csv(io.BytesIO(content), encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError("Could not decode CSV with any supported encoding")
        
        # Build text representation
        text_parts = [f"CSV File: {filename}\n"]
        text_parts.append(f"Columns: {', '.join(df.columns.tolist())}\n")
        text_parts.append(f"Rows: {len(df)}\n\n")
        
        # Add column info
        text_parts.append("Column Details:\n")
        for col in df.columns:
            text_parts.append(f"- {col}: {df[col].dtype}\n")
        text_parts.append("\n")
        
        # Add sample data (first 10 rows)
        text_parts.append("Sample Data:\n")
        text_parts.append(df.head(10).to_string(index=False))
        
        full_text = "".join(text_parts)
        metadata = {
            "columns": df.columns.tolist(),
            "rows": len(df),
            "shape": df.shape,
            "type": "csv"
        }
        
        return full_text, metadata
        
    except Exception as e:
        return f"Error parsing CSV: {str(e)}", {"error": str(e), "type": "csv"}

def _parse_xlsx_content(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
    """Parse Excel file using openpyxl and pandas"""
    try:
        # Load workbook to get sheet names
        wb = load_workbook(io.BytesIO(content), read_only=True)
        sheet_names = wb.sheetnames
        wb.close()
        
        text_parts = [f"Excel File: {filename}\n"]
        text_parts.append(f"Sheets: {', '.join(sheet_names)}\n\n")
        
        # Parse each sheet with pandas
        excel_file = pd.ExcelFile(io.BytesIO(content))
        
        for sheet_name in sheet_names:
            try:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                text_parts.append(f"--- Sheet: {sheet_name} ---\n")
                text_parts.append(f"Columns: {', '.join(df.columns.tolist())}\n")
                text_parts.append(f"Rows: {len(df)}\n")
                
                # Add sample data (first 5 rows per sheet)
                if not df.empty:
                    text_parts.append("Sample Data:\n")
                    text_parts.append(df.head(5).to_string(index=False))
                text_parts.append("\n\n")
                
            except Exception as sheet_error:
                text_parts.append(f"--- Sheet: {sheet_name} (Error) ---\n")
                text_parts.append(f"Could not parse sheet: {str(sheet_error)}\n\n")
        
        full_text = "".join(text_parts)
        metadata = {
            "sheets": sheet_names,
            "sheet_count": len(sheet_names),
            "type": "xlsx"
        }
        
        return full_text, metadata
        
    except Exception as e:
        return f"Error parsing Excel file: {str(e)}", {"error": str(e), "type": "xlsx"}

def _parse_docx_content(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
    """Parse Word document using python-docx"""
    try:
        doc = Document(io.BytesIO(content))
        
        text_parts = [f"Word Document: {filename}\n\n"]
        
        # Extract paragraphs
        paragraph_count = 0
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                text_parts.append("\n")
                paragraph_count += 1
        
        # Extract tables
        table_count = 0
        for table in doc.tables:
            text_parts.append(f"\n--- Table {table_count + 1} ---\n")
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                text_parts.append(row_text)
                text_parts.append("\n")
            table_count += 1
            text_parts.append("\n")
        
        full_text = "".join(text_parts)
        metadata = {
            "paragraphs": paragraph_count,
            "tables": table_count,
            "type": "docx"
        }
        
        return full_text, metadata
        
    except Exception as e:
        return f"Error parsing Word document: {str(e)}", {"error": str(e), "type": "docx"}

# Legacy FileParser class for backward compatibility
class FileParser:
    
    @staticmethod
    def detect_file_type(filename: str, content_type: str) -> str:
        """Legacy method - use detect_file_type function instead"""
        return detect_file_type(filename, content_type)
    
    @staticmethod
    async def parse_file(file_content: bytes, file_type: str, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Legacy method - use parse_file_by_type function instead"""
        result = parse_file_by_type(filename, file_content)
        return result["text"], {"type": file_type, "preview": result["preview"]} 